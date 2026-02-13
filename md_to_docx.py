#!/usr/bin/env python3
"""
Convert Markdown to DOCX with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted DOCX"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lang = None
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start code block
                in_code_block = True
                code_lang = line[3:].strip() or 'text'
                code_lines = []
            else:
                # End code block
                in_code_block = False
                if code_lines:
                    # Add code block to document
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)

                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)

                    # Add background color effect with shading
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.size = Pt(24)
                p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            elif level == 3:
                p = doc.add_heading(text, level=3)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            else:
                p = doc.add_paragraph(text)
                p.runs[0].font.bold = True

            i += 1
            continue

        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph('_' * 80)
            p.runs[0].font.color.rgb = RGBColor(192, 192, 192)
            i += 1
            continue

        # Lists
        if re.match(r'^[\*\-\+]\s+', line) or re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^[\*\-\+\d\.]\s+', '', line).strip()
            p = doc.add_paragraph(text, style='List Bullet' if line.strip()[0] in '*-+' else 'List Number')
            format_inline_styles(p)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j].rstrip())
                j += 1

            # Skip separator line if present
            if len(table_lines) > 1 and re.match(r'^\|[\s\-\|:]+\|$', table_lines[1]):
                table_lines.pop(1)

            # Create table
            rows = []
            for tline in table_lines:
                cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                rows.append(cells)

            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        if row_idx == 0:  # Header row
                            cell.paragraphs[0].runs[0].font.bold = True

            i = j
            continue

        # Checkboxes
        if re.match(r'^-\s+\[[ x]\]\s+', line):
            checked = '[x]' in line
            text = re.sub(r'^-\s+\[[ x]\]\s+', '', line).strip()
            p = doc.add_paragraph()
            p.add_run('☑ ' if checked else '☐ ')
            p.add_run(text)
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph(line)
        format_inline_styles(p)
        i += 1

    # Save document
    doc.save(docx_file)
    print(f"[SUCCESS] Converted {md_file} to {docx_file}")

def format_inline_styles(paragraph):
    """Format inline markdown styles (bold, italic, code)"""
    text = paragraph.text
    paragraph.clear()

    # Simple pattern matching for inline styles
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|_.*?_)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    import sys

    md_file = 'windows-ssh-setup.md'
    docx_file = 'Windows_SSH_Setup_Guide.docx'

    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]

    parse_markdown_to_docx(md_file, docx_file)
